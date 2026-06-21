# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── 全局样式 ──
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.line_spacing = 1.35
style.paragraph_format.space_after = Pt(4)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = '微软雅黑'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    hs.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    if level == 1:
        hs.paragraph_format.space_before = Pt(18)
        hs.paragraph_format.space_after = Pt(10)
    else:
        hs.paragraph_format.space_before = Pt(12)
        hs.paragraph_format.space_after = Pt(6)

code_style = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
code_style.font.name = 'Consolas'
code_style.font.size = Pt(8)
code_style.paragraph_format.space_before = Pt(2)
code_style.paragraph_format.space_after = Pt(2)
code_style.paragraph_format.line_spacing = 1.0

def add_code_block(code_text):
    for line in code_text.strip().split('\n'):
        doc.add_paragraph(line, style='CodeBlock')

def add_table(headers, rows):
    t = doc.add_table(rows=len(rows)+1, cols=len(headers), style='Light Grid Accent 1')
    for j, h in enumerate(headers):
        t.rows[0].cells[j].text = h
    for i, row in enumerate(rows):
        for j, c in enumerate(row):
            t.rows[i+1].cells[j].text = c
    return t

# ═══════════════════════ 封面 ═══════════════════════
for _ in range(5):
    doc.add_paragraph('')

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('四旋翼飞行器飞控系统\n动态角度误差测试与达标计划')
run.font.size = Pt(24)
run.font.bold = True
run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

doc.add_paragraph('')
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('目标：±2° 动态角度误差\n基于 STM32F103C8 + FreeRTOS + MPU6050 + Mahony 滤波')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

for _ in range(3):
    doc.add_paragraph('')

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('文档状态：执行计划\n日    期：2026-05-29')
run.font.size = Pt(11)

doc.add_page_break()

# ═══════════════════════ 概述 ═══════════════════════
doc.add_heading('1. 目标概述', level=1)

doc.add_paragraph('目标是使飞控系统在动态飞行（非静态悬停）中的角度跟踪误差控制在 ±2° 以内。')
doc.add_paragraph('动态误差定义为：在打杆机动期间，遥控目标角度与 Mahony 滤波器输出的实际角度之间的偏差。')

doc.add_heading('1.1 可行性分析', level=2)
doc.add_paragraph(
    'MPU6050 陀螺噪声密度约 0.05°/s/√Hz，经 44Hz DLPF 后噪声约 ±0.3°/s，'
    '控制周期 5ms 内角度误差累积仅约 0.002°。传感器的物理噪声量级远小于 2° 目标。'
)
doc.add_paragraph('误差的主要来源是软件层面：')
doc.add_paragraph('PID 全部 Ki=0，缺乏积分项消除稳态误差，角度永远差一个固定偏差追不上目标。', style='List Bullet')
doc.add_paragraph('Mahony 滤波器 Ki=0，陀螺零偏不被在线估计纠正，长时间积分产生漂移。', style='List Bullet')
doc.add_paragraph('固定权重互补滤波（已替换为 Mahony）在动态场景下误差累积更快。', style='List Bullet')
doc.add_paragraph('以上问题均可通过软件修改解决，±2° 在 F103 平台上完全可达。')

# ═══════════════════════ 代码准备 ═══════════════════════
doc.add_heading('2. 代码准备（一次性工作，约 15 分钟）', level=1)

add_table(
    ('序号', '事项', '涉及文件', '说明'),
    [
        ('1', '扩展 LogEntry_t', 'logger.h', '增加 target_roll / target_pitch / target_yaw_rate 三个 float 字段'),
        ('2', '更新 LOG_ENTRY_SIZE', 'logger.h', '48 → 60（+12 字节）'),
        ('3', '填入目标值', 'main.c', 'Fly_Control_Task 中把 Fly_target 的目标值抄入 log_entry'),
        ('4', '提高日志频率', 'main.c', 'log_tick % 10 → % 2，从 50Hz 提到 100Hz'),
        ('5', '给内环 PID 加 Ki', 'fly_ctrl.c', 'roll_rate / pitch_rate Ki=0.12, yaw_rate Ki=0.08'),
        ('6', '给 Mahony 开 Ki', 'main.c', 'MPU6050_Init 后调用 MPU6050_SetMahonyKi(0.01f)'),
    ]
)

doc.add_paragraph('')
doc.add_paragraph('修改要点（fly_ctrl.c: Fly_Init）：')
add_code_block("""PID_Init(&pid_roll_rate,   0.6f, 0.12f, 0.06f, 50.0f, 300.0f);  // Ki: 0 → 0.12
PID_Init(&pid_pitch_rate,  0.6f, 0.12f, 0.06f, 50.0f, 300.0f);
PID_Init(&pid_yaw_rate,    1.0f, 0.08f, 0.00f, 50.0f, 300.0f);  // Ki: 0 → 0.08""")

doc.add_heading('2.1 为什么 Ki 是关键', level=3)
doc.add_paragraph(
    'PID 控制器中，Kp 和 Kd 都是瞬时信号驱动。飞机受重力、摩擦、电机不对称等干扰，'
    '永远存在一个微小的"残留偏差"使 Kp·error 恰好抵消干扰，但此时 error ≠ 0。'
    '只有 Ki 能通过累积过去的误差来消除这个残留。没有 Ki，稳态误差就不可能为零。'
)

# ═══════════════════════ 测试台 ═══════════════════════
doc.add_heading('3. 测试台搭建（约 1 小时）', level=1)

add_table(
    ('序号', '事项', '说明'),
    [
        ('1', '单轴横滚转台', '滚珠轴承 + 铝型材支架，飞机刚性固定，只能绕横滚轴自由旋转'),
        ('2', '连接电机和桨', '建议先不装桨测一版，确认数据正常后再装桨'),
        ('3', '上电确认', '串口有 printf 输出、SD 卡初始化成功、传感器数据正常'),
    ]
)

doc.add_paragraph('')
doc.add_paragraph(
    '为什么要用转台不飞：飞起来变量太多（机身振动、阵风、操作手抖）。'
    '转台只保留一个旋转轴，所有误差来源除了传感器就是控制算法，结果干净可比。'
)

# ═══════════════════════ 测试操作 ═══════════════════════
doc.add_heading('4. 测试操作（每轮约 10 分钟）', level=1)

doc.add_heading('4.0 准备工作', level=2)
doc.add_paragraph('遥控器设定：')
doc.add_paragraph('横滚、俯仰、偏航通道的 EXP（指数曲线）全部关闭，摇杆响应设为线性。', style='List Bullet')
doc.add_paragraph('横滚、俯仰的舵量（EPA/End Point）设为 100%。', style='List Bullet')
doc.add_paragraph('油门通道保持正常，但测试时保持悬停油门不变（打杆只动横滚/俯仰）。', style='List Bullet')
doc.add_paragraph('')
doc.add_paragraph('上电后观察：')
doc.add_paragraph('SD 卡初始化打印 "SD card ready" → 确认日志在写。', style='List Bullet')
doc.add_paragraph('串口 Monitor_Task 每 200ms 输出角度数据 → 确认传感器正常。', style='List Bullet')
doc.add_paragraph('飞机放在转台上，确认转台能自由绕横滚轴旋转（无卡涩）。', style='List Bullet')
doc.add_paragraph('武装（Armed）成功后，油门推到悬停位置（约 50%）。', style='List Bullet')

doc.add_heading('4.1 第1轮：静态基线（2 分钟）', level=2)
doc.add_paragraph('目的：测出转台上无控状态下的角度漂移量。')
doc.add_paragraph('操作步骤：')
doc.add_paragraph('① 武装飞机，油门推到悬停位，整个飞行器在转台上保持静止。', style='List Bullet')
doc.add_paragraph('② 双手不碰遥控杆，保持 30 秒。', style='List Bullet')
doc.add_paragraph('③ 降落 → 断电 → 拔 SD 卡读日志。', style='List Bullet')
doc.add_paragraph('判断标准：30 秒内角度漂移 < 1°，说明 Mahony 滤波器基线正常。')
doc.add_paragraph('如果漂移 > 3°，说明 Mahony 的 Ki 没开或者零偏校准不够。')

doc.add_heading('4.2 第2轮：小阶跃（3 分钟）', level=2)
doc.add_paragraph('目的：测 10° 小角度下 PID 跟踪精度和稳态误差。')
doc.add_paragraph('操作步骤：')
doc.add_paragraph('① 悬停稳定后，右手握摇杆，准备打横滚。', style='List Bullet')
doc.add_paragraph('② 突然把横滚杆向右打到约 1/3 行程（对应目标 10°），保持 3 秒不动。', style='List Bullet')
doc.add_paragraph('③ 突然回中，保持 2 秒。', style='List Bullet')
doc.add_paragraph('④ 重复 ②~③ 共 3 次。', style='List Bullet')
doc.add_paragraph('⑤ 降落 → 断电 → 读日志。', style='List Bullet')
doc.add_paragraph('关键观察：打杆瞬间角度是直接跟上，还是先反冲一下再追？')
doc.add_paragraph('回中后角度是回到 0°，还是停在 1°~2°？如果回不到 0° → Ki 不够。')

doc.add_heading('4.3 第3轮：大阶跃（3 分钟）', level=2)
doc.add_paragraph('目的：测 30° 大角度下的超调和响应时间。')
doc.add_paragraph('操作步骤：')
doc.add_paragraph('① 悬停稳定后，右手握摇杆。', style='List Bullet')
doc.add_paragraph('② 突然把横滚杆向右打到最大行程（对应目标 30°），保持 3 秒不动。', style='List Bullet')
doc.add_paragraph('③ 突然回中，保持 3 秒。', style='List Bullet')
doc.add_paragraph('④ 重复 ②~③ 共 3 次。', style='List Bullet')
doc.add_paragraph('⑤ 降落 → 断电 → 读日志。', style='List Bullet')
doc.add_paragraph('关键观察：打杆后角度是平稳逼近 30°，还是先冲到 40° 再回调（超调）？')
doc.add_paragraph('超调量 = (峰值角度 - 30°) / 30° × 100%。超过 20% → 需要减小 Ki 或增大 Kd。')

doc.add_heading('4.4 第4轮：正弦扫频（2 分钟）', level=2)
doc.add_paragraph('目的：测飞控的频率响应截止点，找出哪个频率开始跟丢。')
doc.add_paragraph('操作步骤：')
doc.add_paragraph('① 悬停稳定后，右手横滚杆。', style='List Bullet')
doc.add_paragraph('② 以约 0.5Hz（2 秒一个来回）小幅度来回摇动摇杆，幅度约 15°（半杆）。', style='List Bullet')
doc.add_paragraph('③ 逐渐加快到 1Hz、2Hz、3Hz（最快 0.33 秒一个来回），每个频率晃 5 秒。', style='List Bullet')
doc.add_paragraph('④ 回中 → 降落 → 断电 → 读日志。', style='List Bullet')
doc.add_paragraph('判断标准：从日志绘制 target_roll 和 roll 两条曲线，看从哪个频率开始两条线错开、幅度缩小。')
doc.add_paragraph('如果 2Hz 以上角度幅度只有目标值的 50%，说明控制带宽不够。')

doc.add_heading('4.5 第5轮：俯仰验证（10 分钟）', level=2)
doc.add_paragraph('目的：确认横滚方向调好的参数在俯仰方向也有效。')
doc.add_paragraph('操作步骤：')
doc.add_paragraph('① 把飞机从转台上拆下，换个方向重新固定（或者直接飞）。', style='List Bullet')
doc.add_paragraph('② 重复第 2~4 轮全部操作，方向换成俯仰。', style='List Bullet')
doc.add_paragraph('③ 对比两轴的 RMS 误差，差异应 < 10%。', style='List Bullet')

doc.add_paragraph('')
doc.add_paragraph('每轮连续测 3 次取平均，避免单次偶然。每轮之间断电重启，确保积分项清零。')
doc.add_paragraph('所有操作建议录像，便于后续对照日志回放。')

# ═══════════════════════ 日志提取 ═══════════════════════
doc.add_heading('5. 日志提取', level=1)

add_table(
    ('方式', '操作', '适用场景'),
    [
        ('SD 卡直接读', '拔卡 → 读卡器 → Python 解析扇区', '常规：数据完整可靠'),
        ('UART 逐扇区 dump', '写一个 CLI 命令 logdump，逐扇区以 hex 格式输出到串口', '备用：不便拔卡时'),
    ]
)

doc.add_paragraph('')
doc.add_paragraph('SD 卡上的数据格式：每扇区 512 字节，每条记录 60 字节，每扇区 8 条记录。')
doc.add_paragraph('记录结构（按偏移）：')
add_code_block("""偏移   字段             类型       说明
0~3    magic           uint32_t   0x4C4F4745
4~7    seq             uint32_t   序列号
8~11   timestamp_ms    uint32_t   系统时间(ms)
12~15  roll            float      实际横滚角(°)
16~19  pitch           float      实际俯仰角(°)
20~23  yaw             float      实际偏航角(°)
24~27  gyro_x          float      陀螺X(°/s)
28~31  gyro_y          float      陀螺Y(°/s)
32~35  gyro_z          float      陀螺Z(°/s)
36~39  m1              float      电机1(μs)
40~43  m2              float      电机2(μs)
44~47  m3              float      电机3(μs)
48~51  m4              float      电机4(μs)
52~55  target_roll     float      目标横滚(°)   ← 新增
56~59  target_pitch    float      目标俯仰(°)   ← 新增
60~63  target_yaw_rate float      目标偏航(°/s)  ← 新增""")

# ═══════════════════════ 数据分析 ═══════════════════════
doc.add_heading('6. Python 数据分析（约 10 分钟 / 轮）', level=1)

doc.add_paragraph('读 SD 原始扇区、解析并计算指标：')
add_code_block("""import struct, math
import matplotlib.pyplot as plt

def parse_log(bin_path):
    entries = []
    with open(bin_path, 'rb') as f:
        while chunk := f.read(512):
            for i in range(512 // 64):  # 60 字节记录，保守用 64
                off = i * 64
                magic = struct.unpack_from('<I', chunk, off)[0]
                if magic != 0x4C4F4745: break
                vals = struct.unpack_from('<III9f4f3f', chunk, off)
                entries.append({
                    'ts': vals[2],         'roll': vals[3],
                    'pitch': vals[4],      'yaw': vals[5],
                    'target_roll': vals[15], 'target_pitch': vals[16],
                    'target_yaw': vals[17],
                    'target_yaw_rate': vals[18], 'm1': vals[12],
                    'm2': vals[13], 'm3': vals[14], 'm4': vals[15],
                })
    return entries

entries = parse_log('sd_dump.bin')

# 只取打杆期间（target > 5° 判定为机动）
maneuver = [e for e in entries if abs(e['target_roll']) > 5.0]
errors = [abs(e['target_roll'] - e['roll']) for e in maneuver]

rms = math.sqrt(sum(e**2 for e in errors) / len(errors))
max_err = max(errors)
print(f'RMS 动态误差: {rms:.2f}°')
print(f'最大瞬时偏差: {max_err:.2f}°')""")

doc.add_paragraph('')
doc.add_heading('6.1 指标与期望值', level=2)

add_table(
    ('指标', '计算方式', '合格线', '良好线'),
    [
        ('RMS 动态误差', 'sqrt(avg((target-actual)²)), 仅机动期', '< 2.0°', '< 1.0°'),
        ('最大瞬时偏差', 'max(|target-actual|), 阶跃后 0.5s 内', '< 3.0°', '< 1.5°'),
        ('稳态误差', '阶跃后 1~3 秒的平均偏差', '< 0.5°', '< 0.2°'),
        ('超调量', '(overshoot / target) × 100%', '< 20%', '< 10%'),
        ('调节时间', '角度进入目标±1° 的时间', '< 100ms', '< 50ms'),
    ]
)

# ═══════════════════════ 调参 ═══════════════════════
doc.add_heading('7. 调参迭代流程', level=1)

doc.add_paragraph('在飞行中通过 UART CLI 实时调参，无需重新编译：')

add_table(
    ('操作', 'CLI 命令', '如果过冲', '如果反应慢'),
    [
        ('调 Ki', 'pid set roll_rate 0.6 0.15 0.06', '减小 Ki（-0.02）', '增大 Ki（+0.02）'),
        ('调 Kd', 'pid set roll_rate 0.6 0.12 0.10', '增大 Kd（+0.02）', '减小 Kd（-0.02）'),
        ('调 Kp', 'pid set roll_rate 0.8 0.12 0.06', '减小 Kp（-0.1）', '增大 Kp（+0.1）'),
    ]
)

doc.add_paragraph('')
doc.add_paragraph('调参原则：')
doc.add_paragraph('先只调 Ki：从 0.05 开始，每次 +0.05，直到出现轻微超调，再回退 0.02。', style='List Bullet')
doc.add_paragraph('如果仍有振荡但不超调，增大 Kd：每次 +0.02。', style='List Bullet')
doc.add_paragraph('如果响应明显滞后于目标，增大 Kp：每次 +0.1，但不超过 1.5。', style='List Bullet')
doc.add_paragraph('每改一次参数，重新走一遍第 4 节全部 5 轮测试，对比 RMS 误差变化。', style='List Bullet')

doc.add_paragraph('')
add_table(
    ('轮次', '预期 RMS 误差', '说明'),
    [
        ('Ki=0.00 基准测试', '3°~5°', '现有代码，全零 Ki'),
        ('Ki=0.08 第 1 轮', '2°~3°', '积分项开始发挥作用'),
        ('Ki=0.12 第 2 轮', '1.5°~2.5°', '接近目标'),
        ('Ki=0.15 第 3 轮', '1°~2°', '如果不过冲，可以到这一级'),
        ('最终微调', '< 2°', 'Kp/Kd 微调，达标'),
    ]
)

# ═══════════════════════ 验收 ═══════════════════════
doc.add_heading('8. 验收标准', level=1)

doc.add_paragraph('连续 3 次大阶跃测试（横滚 30°、俯仰 30°），同时满足以下全部条件即为达标：')
doc.add_paragraph('RMS 动态误差 ≤ 2.0°', style='List Bullet')
doc.add_paragraph('最大瞬时偏差 ≤ 3.5°', style='List Bullet')
doc.add_paragraph('稳态误差 ≤ 0.5°', style='List Bullet')
doc.add_paragraph('超调量 < 25%', style='List Bullet')

doc.add_paragraph('')
doc.add_paragraph('不合格 → 继续第 7 节调参。')
doc.add_paragraph('合格 → 记录最终 PID 参数和对应的 RMS 误差，保存日志文件留档。')

# ═══════════════════════ 时间估算 ═══════════════════════
doc.add_heading('9. 时间估算', level=1)

add_table(
    ('阶段', '耗时', '说明'),
    [
        ('代码准备', '15 分钟', '一次性工作'),
        ('测试台搭建', '1 小时', '一次性工作'),
        ('每轮飞测', '10 分钟', '需重复 3~5 轮'),
        ('Python 分析', '10 分钟 / 轮', '每轮后分析'),
        ('总预计', '2~3 小时', '含 3 轮调参迭代'),
    ]
)

# 保存
output_path = os.path.join(os.path.dirname(__file__), '后续计划.docx')
doc.save(output_path)
print(f'文档已保存至: {output_path}')
